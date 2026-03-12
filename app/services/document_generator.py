import json
import os
from datetime import datetime
from uuid import uuid4

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.models.enums import DOCUMENT_TYPE_LABELS
from app.prompts.document_prompts import DOCUMENT_SYSTEM_PROMPT, SECTION_PROMPTS
from app.services.knowledge import KnowledgeBase
from app.services.llm import LLMService
from config import settings


class DocumentGenerator:
    def __init__(self, llm: LLMService, knowledge: KnowledgeBase):
        self.llm = llm
        self.knowledge = knowledge
        self.output_dir = settings.GENERATED_DOCS_DIR

    def generate(self, doc_type: str, case_info: dict) -> tuple[str, str]:
        doc = Document()
        self._set_default_styles(doc)
        self._set_page_margins(doc)

        method = getattr(self, f"_generate_{doc_type}", None)
        if method:
            method(doc, case_info)
        else:
            self._generate_generic(doc, doc_type, case_info)

        file_id = uuid4().hex[:12]
        filename = f"{file_id}.docx"
        filepath = os.path.join(self.output_dir, filename)
        os.makedirs(self.output_dir, exist_ok=True)
        doc.save(filepath)
        return file_id, filepath

    def _set_default_styles(self, doc: Document):
        style = doc.styles["Normal"]
        font = style.font
        font.name = "仿宋"
        font.size = Pt(14)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        pf = style.paragraph_format
        pf.line_spacing = 1.5
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

    def _set_page_margins(self, doc: Document):
        for section in doc.sections:
            section.top_margin = Cm(3.7)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)

    def _add_title(self, doc: Document, text: str):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = "宋体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        p.paragraph_format.space_after = Pt(20)

    def _add_heading(self, doc: Document, text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.name = "黑体"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

    def _add_paragraph(self, doc: Document, text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.name = "仿宋"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
        p.paragraph_format.first_line_indent = Pt(28)

    def _add_field(self, doc: Document, label: str, value: str):
        p = doc.add_paragraph()
        run = p.add_run(f"{label}：{value}")
        run.font.size = Pt(14)
        run.font.name = "仿宋"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    def _add_closing(self, doc: Document, recipient: str, name: str):
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(28)
        run = p.add_run(f"此致")
        run.font.size = Pt(14)
        run.font.name = "仿宋"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run2 = p2.add_run(recipient)
        run2.font.size = Pt(14)
        run2.font.name = "仿宋"
        run2.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

        doc.add_paragraph()

        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run3 = p3.add_run(f"投诉人：{name}")
        run3.font.size = Pt(14)
        run3.font.name = "仿宋"
        run3.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        today = datetime.now().strftime("%Y年%m月%d日")
        run4 = p4.add_run(today)
        run4.font.size = Pt(14)
        run4.font.name = "仿宋"
        run4.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    def _call_llm_for_section(self, prompt_key: str, case_info: dict) -> str:
        relevant_laws = self.knowledge.get_law_text_for_prompt(case_info.get("case_type"))
        case_info_text = json.dumps(
            {k: v for k, v in case_info.items() if not k.startswith("_")},
            ensure_ascii=False,
            indent=2,
        )
        prompt = SECTION_PROMPTS[prompt_key].format(
            case_info=case_info_text,
            relevant_laws=relevant_laws,
        )
        messages = [
            {"role": "system", "content": DOCUMENT_SYSTEM_PROMPT},
            {"role": "user", "content": prompt},
        ]
        return self.llm.chat(messages)

    # --- Specific document generators ---

    def _generate_complaint_letter(self, doc: Document, case_info: dict):
        self._add_title(doc, "投 诉 书")

        self._add_heading(doc, "一、投诉人信息")
        self._add_field(doc, "姓名", case_info.get("complainant_name", "________"))
        self._add_field(doc, "联系电话", case_info.get("complainant_phone", "________"))
        self._add_field(doc, "地址", case_info.get("complainant_address", "________"))

        self._add_heading(doc, "二、被投诉人信息")
        self._add_field(doc, "名称", case_info.get("merchant_name", "________"))
        self._add_field(doc, "地址", case_info.get("merchant_address", "________"))

        self._add_heading(doc, "三、投诉请求")
        demands = self._call_llm_for_section("demands_section", case_info)
        self._add_paragraph(doc, demands)

        self._add_heading(doc, "四、事实与理由")
        facts = self._call_llm_for_section("facts_and_reasons", case_info)
        self._add_paragraph(doc, facts)

        self._add_heading(doc, "五、证据清单")
        evidence = self._call_llm_for_section("evidence_analysis", case_info)
        self._add_paragraph(doc, evidence)

        self._add_closing(
            doc,
            "市场监督管理局",
            case_info.get("complainant_name", "________"),
        )

    def _generate_report_letter(self, doc: Document, case_info: dict):
        self._add_title(doc, "举 报 信")

        self._add_heading(doc, "一、举报人信息")
        self._add_field(doc, "姓名", case_info.get("complainant_name", "________"))
        self._add_field(doc, "联系电话", case_info.get("complainant_phone", "________"))

        self._add_heading(doc, "二、被举报人信息")
        self._add_field(doc, "名称", case_info.get("merchant_name", "________"))
        self._add_field(doc, "地址", case_info.get("merchant_address", "________"))

        self._add_heading(doc, "三、举报事项")
        content = self._call_llm_for_section("report_content", case_info)
        self._add_paragraph(doc, content)

        self._add_heading(doc, "四、证据材料")
        evidence = self._call_llm_for_section("evidence_analysis", case_info)
        self._add_paragraph(doc, evidence)

        self._add_closing(
            doc,
            "市场监督管理局",
            case_info.get("complainant_name", "________"),
        )

    def _generate_demand_letter(self, doc: Document, case_info: dict):
        self._add_title(doc, "协 商 函")

        merchant = case_info.get("merchant_name", "________")
        p = doc.add_paragraph()
        run = p.add_run(f"{merchant}：")
        run.font.size = Pt(14)
        run.font.name = "仿宋"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

        content = self._call_llm_for_section("demand_content", case_info)
        self._add_paragraph(doc, content)

        self._add_closing(
            doc,
            merchant,
            case_info.get("complainant_name", "________"),
        )

    def _generate_civil_lawsuit(self, doc: Document, case_info: dict):
        self._add_title(doc, "民 事 起 诉 状")

        self._add_heading(doc, "原告")
        self._add_field(doc, "姓名", case_info.get("complainant_name", "________"))
        self._add_field(doc, "身份证号", case_info.get("complainant_id_number", "________"))
        self._add_field(doc, "联系电话", case_info.get("complainant_phone", "________"))
        self._add_field(doc, "住址", case_info.get("complainant_address", "________"))

        self._add_heading(doc, "被告")
        self._add_field(doc, "名称", case_info.get("merchant_name", "________"))
        self._add_field(doc, "地址", case_info.get("merchant_address", "________"))

        self._add_heading(doc, "诉讼请求")
        claims = self._call_llm_for_section("lawsuit_claims", case_info)
        self._add_paragraph(doc, claims)

        self._add_heading(doc, "事实与理由")
        facts = self._call_llm_for_section("lawsuit_facts", case_info)
        self._add_paragraph(doc, facts)

        self._add_heading(doc, "证据清单")
        evidence = self._call_llm_for_section("evidence_analysis", case_info)
        self._add_paragraph(doc, evidence)

        self._add_closing(
            doc,
            "________人民法院",
            case_info.get("complainant_name", "________"),
        )

    def _generate_evidence_checklist(self, doc: Document, case_info: dict):
        self._add_title(doc, "证 据 清 单")

        evidence = self._call_llm_for_section("evidence_analysis", case_info)
        self._add_paragraph(doc, evidence)

        doc.add_paragraph()
        today = datetime.now().strftime("%Y年%m月%d日")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"整理日期：{today}")
        run.font.size = Pt(14)
        run.font.name = "仿宋"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    def _generate_claim_letter(self, doc: Document, case_info: dict):
        self._add_title(doc, "索 赔 函")

        merchant = case_info.get("merchant_name", "________")
        p = doc.add_paragraph()
        run = p.add_run(f"{merchant}：")
        run.font.size = Pt(14)
        run.font.name = "仿宋"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

        content = self._call_llm_for_section("claim_content", case_info)
        self._add_paragraph(doc, content)

        self._add_closing(
            doc,
            merchant,
            case_info.get("complainant_name", "________"),
        )

    def _generate_generic(self, doc: Document, doc_type: str, case_info: dict):
        label = DOCUMENT_TYPE_LABELS.get(doc_type, doc_type)
        self._add_title(doc, label)
        facts = self._call_llm_for_section("facts_and_reasons", case_info)
        self._add_paragraph(doc, facts)
